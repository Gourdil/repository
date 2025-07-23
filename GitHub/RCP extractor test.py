import sys
import requests
from bs4 import BeautifulSoup
from docx import Document

def chercher_url_fiche(nom_medicament):
    url_recherche = f"https://base-donnees-publique.medicaments.gouv.fr/recherche.php?searchtype=medicament&txtNom={nom_medicament}"
    res = requests.get(url_recherche, headers={"User-Agent": "Mozilla/5.0"})
    soup = BeautifulSoup(res.text, "html.parser")
    lien = soup.find("a", href=lambda x: x and "fichemedicament" in x)
    if lien:
        return "https://base-donnees-publique.medicaments.gouv.fr" + lien['href']
    return None

def chercher_url_rcp(url_fiche):
    res = requests.get(url_fiche)
    soup = BeautifulSoup(res.text, "html.parser")
    lien_rcp = soup.find("a", href=lambda x: x and "typedoc=R" in x)
    if lien_rcp:
        return "https://base-donnees-publique.medicaments.gouv.fr/" + lien_rcp['href']
    return None

def extraire_sections_rcp(url_rcp, debut="4.1", fin="5.3"):
    res = requests.get(url_rcp)
    soup = BeautifulSoup(res.text, "html.parser")
    sections = soup.select("div.titreRCP, div.texteRCP")

    extractions = {}
    current_title = None
    recording = False

    for elem in sections:
        if "titreRCP" in elem.get("class", []):
            titre = elem.get_text(strip=True)
            if titre.startswith(debut):
                recording = True
            elif titre.startswith("5.4") or titre.startswith("6."):
                recording = False
            if recording:
                current_title = titre
                extractions[current_title] = ""
        elif "texteRCP" in elem.get("class", []) and recording and current_title:
            extractions[current_title] += elem.get_text(strip=True, separator="\n") + "\n"

    return extractions

def generer_word(nom_medicament, sections):
    doc = Document()
    doc.add_heading(f"RCP – {nom_medicament}", 0)
    for titre, texte in sections.items():
        doc.add_heading(titre, level=1)
        doc.add_paragraph(texte)
    nom_fichier = f"RCP_{nom_medicament.replace(' ', '_')}.docx"
    doc.save(nom_fichier)
    print(f"✅ Document Word généré : {nom_fichier}")

def main():
    if len(sys.argv) < 2:
        print("Usage : python rcp_extractor.py 'Nom du médicament'")
        return

    nom_med = sys.argv[1]
    print(f"🔎 Recherche de {nom_med}...")

    url_fiche = chercher_url_fiche(nom_med)
    if not url_fiche:
        print("❌ Fiche médicament introuvable.")
        return

    url_rcp = chercher_url_rcp(url_fiche)
    if not url_rcp:
        print("❌ RCP non trouvé.")
        return

    print("📥 Téléchargement et extraction du RCP...")
    sections = extraire_sections_rcp(url_rcp)

    if not sections:
        print("❌ Aucune section 4.1 à 5.3 trouvée.")
        return

    generer_word(nom_med, sections)

if __name__ == "__main__":
    main()